# app.py  – CAAT Avanzado (incluye CxC vs Bancos + Aging)
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import io, csv, re
from datetime import timedelta
from docx import Document
from docx.shared import Pt

# ------------------------- Apariencia -------------------------
st.set_page_config(page_title="CAAT - Conciliación y Auditoría", layout="wide")
st.markdown("""
<style>
.main .block-container {max-width: 1250px; padding-top: 0.6rem; padding-bottom: 2rem;}
.stTabs [data-baseweb="tab-list"] { gap: 10px; }
.stTabs [data-baseweb="tab"]{
  height: 58px; border-radius: 12px !important; padding: 12px 20px !important;
  background: #f6f7fb; border: 1px solid rgba(47,58,178,.15);
  font-size: 17px !important; font-weight: 700;
}
.stTabs [aria-selected="true"]{
  background: #eef2ff !important; color: #2f3ab2 !important; border: 2px solid #2f3ab2 !important;
}
.section-card {
  border: 1px solid rgba(125,125,125,.22);
  border-radius: 16px; padding: 18px 20px; margin: 16px 0 24px 0;
  background: #ffffff; box-shadow: 0 1px 0 rgba(0,0,0,0.04);
}
.section-title { font-size: 26px; font-weight: 800; margin-bottom: 6px; }
.section-desc  { font-size: 16.5px; color:#374151; }
[data-testid="stFileUploader"]{ border-radius:12px; border:1px dashed rgba(125,125,125,.35); padding:18px;}
.stButton>button{ border-radius:999px !important; padding:.6rem 1.1rem; font-weight:700;}
.big-warning { font-size: 16px; line-height: 1.35; }
</style>
""", unsafe_allow_html=True)

st.title("📊 CAAT – Conciliación y Auditoría Automatizada")
st.caption("Soporta **CSV/XLSX/XLS/TXT**. Descargas en **XLSX** y reportes en **DOCX**. "
           "Incluye conciliación **CxC vs Bancos + Aging** con tolerancias, NC/Retenciones y saldos irrisorios.")

# ------------------------- Utilidades comunes -------------------------
PRUEBAS = [
    "1. Transacciones Conciliadas Completas",
    "2. Faltantes en el Destino (Solo en Origen)",
    "3. Inesperadas en el Destino (Solo en Destino)",
    "4. Discrepancias por ID (Monto/Fecha)",
    "5. Duplicados Internos",
    "6. CxC vs Bancos + Aging"  # NUEVA
]
CAMPOS_CLAVE = ["ID_Transaccion", "Fecha", "Monto", "ID_Entidad"]
CAMPOS_ID = ["ID_Transaccion", "ID_Entidad"]

def sniff_delimiter(sample_bytes: bytes):
    try:
        sample = sample_bytes.decode('utf-8', errors='ignore')
        dialect = csv.Sniffer().sniff(sample, delimiters=";,|\t")
        return dialect.delimiter
    except Exception:
        return None

def read_any(file, widget_key="sheet"):
    name = file.name.lower()
    if name.endswith(".csv") or name.endswith(".txt"):
        data = file.read()
        if isinstance(data, bytes):
            delim = sniff_delimiter(data[:4096])
            bio = io.BytesIO(data)
            if delim:
                try: return pd.read_csv(bio, sep=delim, engine="python")
                except Exception: pass
            bio.seek(0)
            try: return pd.read_csv(bio, sep=None, engine="python")
            except Exception:
                bio.seek(0); return pd.read_csv(bio, sep=None, engine="python", encoding="latin-1")
        else:
            return pd.read_csv(io.StringIO(data))
    else:
        # Excel
        try:
            xls = pd.ExcelFile(file)
            sheet = st.selectbox("📄 Hoja de Excel", xls.sheet_names, key=widget_key)
            return pd.read_excel(xls, sheet_name=sheet)
        except Exception:
            return pd.read_excel(file)

def validar_columnas(df, nombre, requeridas):
    faltantes = [col for col in requeridas if col not in df.columns]
    if faltantes:
        st.error(f"❌ El archivo '{nombre}' no contiene las columnas necesarias: {', '.join(faltantes)}")
        return False
    return True

def to_xlsx_bytes(sheets: dict):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        for name, df in sheets.items():
            nm = str(name)[:31]
            (df if isinstance(df, pd.DataFrame) else pd.DataFrame(df)).to_excel(w, index=False, sheet_name=nm)
    return buf.getvalue()

def docx_from_sections(title: str, sections: list[tuple[str, list[str]]]) -> bytes:
    d = Document(); d.add_heading(title, level=1)
    for heading, bullets in sections:
        d.add_heading(heading, level=2)
        for item in bullets:
            p = d.add_paragraph(item, style="List Bullet")
            p.style.font.size = Pt(11)
    bio = io.BytesIO(); d.save(bio); return bio.getvalue()

def coerce_amount(series: pd.Series) -> pd.Series:
    s = series.astype(str).str.replace(r"\.", "", regex=True).str.replace(",", ".", regex=False)
    return pd.to_numeric(s, errors="coerce")

def coerce_date(series: pd.Series) -> pd.Series:
    return pd.to_datetime(series, errors="coerce", dayfirst=True)

# ------------------------- Panel lateral -------------------------
opcion = st.sidebar.selectbox("Selecciona la prueba CAAT", PRUEBAS)

# ------------------------- PRUEBAS 1–5 (tu base existente) -------------------------
conteo_resultados = {
    "Conciliadas": 0,
    "Faltantes en destino": 0,
    "Inesperadas en destino": 0,
    "Discrepancias de valor": 0,
    "Duplicados": 0
}

def generar_conclusion_conteo(conteo):
    conclusion = "🔍 **Conclusión General del Análisis**\n\nDurante la conciliación se identificaron:\n\n"
    if conteo["Faltantes en destino"] > 0:
        conclusion += f"- **{conteo['Faltantes en destino']}** transacciones ausentes en el destino.\n"
    if conteo["Inesperadas en destino"] > 0:
        conclusion += f"- **{conteo['Inesperadas en destino']}** transacciones inesperadas en el destino.\n"
    if conteo["Discrepancias de valor"] > 0:
        conclusion += f"- **{conteo['Discrepancias de valor']}** discrepancias en monto o fecha.\n"
    if conteo["Duplicados"] > 0:
        conclusion += f"- **{conteo['Duplicados']}** registros duplicados.\n"
    if all(conteo[k] == 0 for k in ["Faltantes en destino","Inesperadas en destino","Discrepancias de valor","Duplicados"]):
        conclusion += "- No se detectaron anomalías significativas. Conciliación correcta.\n"
    conclusion += "\n📌 Prioriza los casos con mayor materialidad e impacto."
    return conclusion

def generar_recomendacion(nombre, cantidad, umbral, mensaje_ok, mensaje_alerta):
    if cantidad > umbral:
        return f"🔴 Riesgo alto en **{nombre}**: {mensaje_alerta}"
    elif cantidad > 0:
        return f"🟡 Atención en **{nombre}**: {mensaje_alerta}"
    else:
        return f"🟢 **{nombre}** en buen estado: {mensaje_ok}"

if opcion != PRUEBAS[4] and opcion != PRUEBAS[5]:
    file_origen = st.file_uploader("📂 Archivo de Origen", type=["xlsx", "xls", "csv", "txt"], key="origen")
    file_destino = st.file_uploader("📁 Archivo de Destino", type=["xlsx", "xls", "csv", "txt"], key="destino")
elif opcion == PRUEBAS[4]:
    file_data = st.file_uploader("📥 Archivo a Analizar", type=["xlsx", "xls", "csv", "txt"], key="uno")

if opcion == PRUEBAS[0] and file_origen and file_destino:
    df1, df2 = read_any(file_origen, "sheet_o"), read_any(file_destino, "sheet_d")
    if validar_columnas(df1, "origen", CAMPOS_CLAVE) and validar_columnas(df2, "destino", CAMPOS_CLAVE):
        df1["Fecha"] = coerce_date(df1["Fecha"]); df2["Fecha"] = coerce_date(df2["Fecha"])
        conciliadas = pd.merge(df1, df2, how="inner", on=CAMPOS_CLAVE)
        conteo_resultados["Conciliadas"] = len(conciliadas)
        st.success(f"✅ {len(conciliadas)} transacciones conciliadas.")
        st.dataframe(conciliadas)
        st.download_button("⬇ Descargar", conciliadas.to_csv(index=False).encode(), "conciliadas.csv", "text/csv")

elif opcion == PRUEBAS[1] and file_origen and file_destino:
    df1, df2 = read_any(file_origen, "sheet_o"), read_any(file_destino, "sheet_d")
    if validar_columnas(df1, "origen", CAMPOS_CLAVE) and validar_columnas(df2, "destino", CAMPOS_CLAVE):
        df1["Fecha"] = coerce_date(df1["Fecha"]); df2["Fecha"] = coerce_date(df2["Fecha"])
        merge = pd.merge(df1, df2, how="left", on=CAMPOS_CLAVE, indicator=True)
        solo_origen = merge[merge["_merge"] == "left_only"].drop(columns="_merge")
        conteo_resultados["Faltantes en destino"] = len(solo_origen)
        st.warning(f"❗ {len(solo_origen)} transacciones solo en el origen.")
        st.dataframe(solo_origen)
        st.download_button("⬇ Descargar", solo_origen.to_csv(index=False).encode(), "solo_origen.csv", "text/csv")

elif opcion == PRUEBAS[2] and file_origen and file_destino:
    df1, df2 = read_any(file_origen, "sheet_o"), read_any(file_destino, "sheet_d")
    if validar_columnas(df1, "origen", CAMPOS_CLAVE) and validar_columnas(df2, "destino", CAMPOS_CLAVE):
        df1["Fecha"] = coerce_date(df1["Fecha"]); df2["Fecha"] = coerce_date(df2["Fecha"])
        merge = pd.merge(df2, df1, how="left", on=CAMPOS_CLAVE, indicator=True)
        solo_destino = merge[merge["_merge"] == "left_only"].drop(columns="_merge")
        conteo_resultados["Inesperadas en destino"] = len(solo_destino)
        st.warning(f"🚨 {len(solo_destino)} transacciones inesperadas en el destino.")
        st.dataframe(solo_destino)
        st.download_button("⬇ Descargar", solo_destino.to_csv(index=False).encode(), "solo_destino.csv", "text/csv")

elif opcion == PRUEBAS[3] and file_origen and file_destino:
    df1, df2 = read_any(file_origen, "sheet_o"), read_any(file_destino, "sheet_d")
    if validar_columnas(df1, "origen", CAMPOS_CLAVE) and validar_columnas(df2, "destino", CAMPOS_CLAVE):
        df1["Fecha"] = coerce_date(df1["Fecha"]); df2["Fecha"] = coerce_date(df2["Fecha"])
        merged = pd.merge(df1, df2, on=CAMPOS_ID, how="inner", suffixes=("_origen", "_destino"))
        discrepancias = merged[(merged["Monto_origen"] != merged["Monto_destino"]) |
                               (merged["Fecha_origen"] != merged["Fecha_destino"])]
        conteo_resultados["Discrepancias de valor"] = len(discrepancias)
        st.warning(f"⚠️ {len(discrepancias)} discrepancias encontradas.")
        st.dataframe(discrepancias)
        st.download_button("⬇ Descargar", discrepancias.to_csv(index=False).encode(), "discrepancias.csv", "text/csv")

elif opcion == PRUEBAS[4] and file_data:
    df = read_any(file_data, "sheet_uno")
    if validar_columnas(df, "archivo único", CAMPOS_CLAVE):
        df["Fecha"] = coerce_date(df["Fecha"])
        duplicados = df[df.duplicated(subset=CAMPOS_CLAVE, keep=False)]
        conteo_resultados["Duplicados"] = len(duplicados)
        st.warning(f"🔁 {len(duplicados)} duplicados encontrados.")
        st.dataframe(duplicados)
        st.download_button("⬇ Descargar", duplicados.to_csv(index=False).encode(), "duplicados.csv", "text/csv")

# ------------------------- NUEVA PRUEBA 6: CxC vs Bancos + Aging -------------------------
if opcion == PRUEBAS[5]:
    st.markdown("""
<div class="section-card">
  <div class="section-title">6️⃣ Cuentas por Cobrar (CxC) vs Bancos + Aging</div>
  <div class="section-desc">
    Concilia saldos de clientes (CxC) con depósitos/transferencias bancarias. Detecta <strong>pagos no aplicados</strong>,
    <strong>NC/retenciones no cruzadas</strong>, <strong>saldos irrisorios</strong> y clasifica antigüedad (0–30/31–60/61–90/90+).
  </div>
</div>
""", unsafe_allow_html=True)

    with st.expander("🧭 ¿Qué descubre esta prueba? / Entregables", expanded=True):
        st.markdown("""
- **Pagos en banco no aplicados** a facturas (Tesorería vs Cobranzas).
- **Facturas vencidas** con alto riesgo (aging configurable).
- **Notas de crédito/retenciones** no aplicadas (indicios por signo y palabras clave).
- **Saldos irrisorios** que deberían sanease según política.
  
**Entregables**  
- **XLSX**: Resumen, Conciliados, PendientesCxC, PagosNoAplicadosBanco, Aging, SaldosIrrisorios, PosiblesNC_Retenciones.  
- **DOCX**: Resumen ejecutivo, hallazgos y **recomendaciones accionables**.
""")

    col1, col2 = st.columns(2)
    with col1:
        file_cxc = st.file_uploader("📂 CxC (facturas/abonos por cliente)", type=["xlsx","xls","csv","txt"], key="cxc")
    with col2:
        file_bank = st.file_uploader("🏦 Extracto bancario", type=["xlsx","xls","csv","txt"], key="bank")

    st.markdown("**Campos mínimos esperados (flexibles en nombre):**")
    st.caption("- CxC: Cliente, NumeroFactura/Referencia, Fecha, Monto (positivo factura, negativo NC/retenciones), Observacion/Glosa (opcional)")
    st.caption("- Banco: Fecha, Monto (depósitos/transferencias +), Referencia/Concepto")

    # Parámetros
    with st.expander("⚙️ Parámetros de conciliación", expanded=True):
        tol_monto = st.number_input("🎯 Tolerancia de monto", min_value=0.0, value=0.50, help="Diferencia máxima para considerar un match.")
        tol_dias = st.number_input("🗓️ Ventana de días entre banco y CxC", min_value=0, value=5, help="Diferencia máxima de fechas para match por monto.")
        irrisorio = st.number_input("🟦 Umbral de saldo irrisorio", min_value=0.0, value=5.0)
        aging_cortes = st.multiselect("📊 Cortes de aging (días)", [30,60,90], default=[30,60,90])
        ejecutar6 = st.button("🔍 Ejecutar conciliación CxC vs Bancos")

    if file_cxc and file_bank and ejecutar6:
        cxc = read_any(file_cxc, "sheet_cxc").rename(columns=lambda x: str(x).strip())
        bank = read_any(file_bank, "sheet_bank").rename(columns=lambda x: str(x).strip())

        # Heurísticas de columnas
        def pick(df, choices):
            for c in df.columns:
                if c.lower() in [s.lower() for s in choices]: return c
                if any(re.search(rf"\b{re.escape(s)}\b", c, flags=re.I) for s in choices): return c
            return None

        col_cli = pick(cxc, ["cliente","id_cliente","ruc","identificacion"])
        col_ref_cxc = pick(cxc, ["numerofactura","numero_factura","referencia","documento","id_transaccion","id"])
        col_fecha_cxc = pick(cxc, ["fecha","fecha_emision","fecha_documento"])
        col_monto_cxc = pick(cxc, ["monto","importe","total","saldo","valor"])
        col_obs_cxc = pick(cxc, ["observacion","glosa","detalle","descripcion"])

        col_fecha_b = pick(bank, ["fecha","fec","date"])
        col_monto_b = pick(bank, ["monto","importe","abono","deposito","cr","credito","valor"])
        col_ref_b = pick(bank, ["referencia","ref","descripcion","concepto","detalle"])

        # Validación mínima
        req_ok = all([col_fecha_cxc, col_monto_cxc, col_fecha_b, col_monto_b])
        if not req_ok:
            st.error("❌ No se pudieron identificar las columnas mínimas (Fecha/Monto) en CxC o Banco.")
            st.stop()

        # Normalizaciones
        cxc["_FECHA"] = coerce_date(cxc[col_fecha_cxc])
        cxc["_MONTO"] = coerce_amount(cxc[col_monto_cxc])
        cxc["_REF"] = cxc[col_ref_cxc].astype(str).str.strip().str.upper() if col_ref_cxc else ""
        cxc["_CLI"] = cxc[col_cli].astype(str).str.strip() if col_cli else "SIN_CLIENTE"
        cxc["_OBS"] = cxc[col_obs_cxc].astype(str) if col_obs_cxc else ""

        bank["_FECHA"] = coerce_date(bank[col_fecha_b])
        bank["_MONTO"] = coerce_amount(bank[col_monto_b])
        bank["_REF"] = bank[col_ref_b].astype(str).str.strip().str.upper() if col_ref_b else ""

        cxc = cxc.dropna(subset=["_FECHA","_MONTO"]).copy()
        bank = bank.dropna(subset=["_FECHA","_MONTO"]).copy()

        # 1) Match exacto por referencia (si hay)
        matched_ref = pd.DataFrame()
        if col_ref_cxc and col_ref_b:
            matched_ref = cxc.merge(bank, on="_REF", suffixes=("_CxC","_Banco"))
            # filtra por tolerancias también (opcional)
            matched_ref = matched_ref[(matched_ref["_MONTO_CxC"] - matched_ref["_MONTO_Banco"]).abs() <= tol_monto]

        # 2) Match por monto (+/- tolerancia) y fecha cercana
        # Redondeo por tolerancia: discretizamos en bandas de tolerancia
        def band_amount(s, tol):
            return (s / max(tol, 0.01)).round(0) if tol > 0 else s.round(2)

        cxc["_BANDA"] = band_amount(cxc["_MONTO"].abs(), tol_monto)
        bank["_BANDA"] = band_amount(bank["_MONTO"].abs(), tol_monto)

        approx = cxc.merge(bank, on="_BANDA", suffixes=("_CxC","_Banco"))
        approx["_DIF_MONTO"] = (approx["_MONTO_CxC"].abs() - approx["_MONTO_Banco"].abs()).abs()
        approx["_DIF_DIAS"] = (approx["_FECHA_CxC"] - approx["_FECHA_Banco"]).abs().dt.days
        approx = approx[(approx["_DIF_MONTO"] <= tol_monto) & (approx["_DIF_DIAS"] <= tol_dias)]

        # Evita duplicidad: prioriza matches por referencia
        if not matched_ref.empty:
            keys_ref = set(zip(matched_ref.index, matched_ref["_REF"]))
            approx = approx[~approx["_REF_CxC"].isin(matched_ref["_REF"])]

        conciliados = pd.concat([matched_ref, approx], ignore_index=True, sort=False).drop_duplicates()

        # 3) Pendientes en CxC (no conciliados) y pagos de banco no aplicados
        idx_cxc_conc = set(conciliados.index)  # NO sirve: cambia index por concat; mejor marcar
        cxc["_ROWID"] = np.arange(len(cxc)); bank["_ROWID"] = np.arange(len(bank))
        if not conciliados.empty:
            if "_ROWID_CxC" not in conciliados.columns and "_ROWID_Banco" not in conciliados.columns:
                # reconstruye rowids por merge alternativo
                conciliados = conciliados.merge(cxc[["_ROWID","_FECHA","_MONTO","_REF","_CLI"]], left_on=["_FECHA_CxC","_MONTO_CxC","_REF_CxC","_CLI_CxC"], right_on=["_FECHA","_MONTO","_REF","_CLI"], how="left")
                conciliados = conciliados.rename(columns={"_ROWID":"_ROWID_CxC"}).drop(columns=["_FECHA","_MONTO","_REF","_CLI"])
                conciliados = conciliados.merge(bank[["_ROWID","_FECHA","_MONTO","_REF"]], left_on=["_FECHA_Banco","_MONTO_Banco","_REF_Banco"], right_on=["_FECHA","_MONTO","_REF"], how="left")
                conciliados = conciliados.rename(columns={"_ROWID":"_ROWID_Banco"}).drop(columns=["_FECHA","_MONTO","_REF"])
        pend_cxc = cxc[~cxc["_ROWID"].isin(conciliados.get("_ROWID_CxC", pd.Series(dtype=int)))].copy()
        pagos_no_aplicados = bank[~bank["_ROWID"].isin(conciliados.get("_ROWID_Banco", pd.Series(dtype=int)))].copy()

        # 4) Aging de pendientes CxC (por días de vencimiento respecto a hoy/fecha más reciente)
        hoy = max(pd.Timestamp.today().normalize(), cxc["_FECHA"].max())
        dias = (hoy - pend_cxc["_FECHA"]).dt.days
        def bucket(d):
            if d <= (aging_cortes[0] if len(aging_cortes)>0 else 30): return "0-30"
            if len(aging_cortes)>1 and d <= aging_cortes[1]: return "31-60"
            if len(aging_cortes)>2 and d <= aging_cortes[2]: return "61-90"
            return "90+"
        if len(pend_cxc):
            pend_cxc["Aging_dias"] = dias
            pend_cxc["Aging_bucket"] = pend_cxc["Aging_dias"].apply(bucket)
        aging = pend_cxc.groupby("Aging_bucket", dropna=False).agg(N=("Aging_dias","count"),
                                                                  Suma=("_MONTO","sum")).reset_index().sort_values("Aging_bucket")

        # 5) Saldos irrisorios
        irrisorios_df = pend_cxc[(pend_cxc["_MONTO"].abs() <= irrisorio)].copy()

        # 6) Heurística NC/Retenciones en CxC (orientativa)
        pos_nc = cxc[(cxc["_MONTO"] < 0) | (cxc["_OBS"].str.contains(r"\b(NC|nota de crédito|retenc)", case=False, na=False))]
        posibles_nc = pos_nc.copy()

        # MÉTRICAS
        c_conc = len(conciliados)
        c_pend = len(pend_cxc)
        c_noap = len(pagos_no_aplicados)
        c_irri = len(irrisorios_df)
        st.subheader("📊 Métricas clave")
        c1,c2,c3,c4 = st.columns(4)
        c1.metric("Conciliados", c_conc)
        c2.metric("Pendientes CxC", c_pend)
        c3.metric("Pagos no aplicados (Banco)", c_noap)
        c4.metric("Saldos irrisorios", c_irri)
        st.caption(f"Ventana ±{tol_dias} días, tolerancia de monto ±{tol_monto:,.2f}")

        with st.expander("🔎 Conciliados", expanded=False): st.dataframe(conciliados.head(1000))
        with st.expander("🟥 Pendientes en CxC", expanded=False): st.dataframe(pend_cxc.head(1000))
        with st.expander("🟧 Pagos en banco no aplicados", expanded=False): st.dataframe(pagos_no_aplicados.head(1000))
        with st.expander("📆 Aging pendientes", expanded=False): st.dataframe(aging)
        if len(irrisorios_df): 
            with st.expander("🟦 Saldos irrisorios", expanded=False): st.dataframe(irrisorios_df.head(1000))
        if len(posibles_nc):
            with st.expander("🟪 Posibles NC/Retenciones", expanded=False): st.dataframe(posibles_nc.head(1000))

        # XLSX
        sheets = {
            "Resumen": pd.DataFrame({
                "Métrica":["Conciliados","PendientesCxC","PagosNoAplicadosBanco","SaldosIrrisorios","VentanaDias","TolMonto"],
                "Valor":[c_conc, c_pend, c_noap, c_irri, tol_dias, tol_monto]
            }),
            "Conciliados": conciliados,
            "PendientesCxC": pend_cxc,
            "PagosNoAplicadosBanco": pagos_no_aplicados,
            "Aging": aging,
            "SaldosIrrisorios": irrisorios_df,
            "PosiblesNC_Retenciones": posibles_nc
        }
        st.download_button("⬇️ Descargar hallazgos CxC vs Bancos (XLSX)",
                           to_xlsx_bytes(sheets), "cxc_bancos_hallazgos.xlsx",
                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # DOCX – recomendaciones
        recs = [
            "Automatizar el cruce de pagos banco ↔ facturas con ventana de días y tolerancia de monto.",
            "Parametrizar aging (0–30/31–60/61–90/90+) con alertas a Cobranzas desde 31+ días.",
            "Saneamiento mensual de saldos irrisorios conforme política (p. ej., ≤ 5 USD).",
            "Forzar aplicación de NC/retenciones contra las facturas correspondientes antes del cierre.",
            "Revisión quincenal conjunta Tesorería–Cobranzas y bitácora de pagos no identificados.",
            "Incluir referencia obligatoria en depósitos (n° factura/cliente) y validar en interfaz bancaria."
        ]
        resumen_doc = [
            f"Archivo CxC: {file_cxc.name} | Banco: {file_bank.name}",
            f"Conciliados: {c_conc} | Pendientes CxC: {c_pend} | Pagos no aplicados: {c_noap}",
            f"Saldos irrisorios (≤ {irrisorio:.2f}): {c_irri}",
            f"Parámetros: Ventana ±{tol_dias} días, tolerancia ±{tol_monto:.2f}"
        ]
        top_focus = []
        if len(aging):
            worst = aging.sort_values("Suma", ascending=False).head(1)
            if len(worst):
                b = worst.iloc[0]["Aging_bucket"]; s = worst.iloc[0]["Suma"]
                top_focus.append(f"Aging crítico: {b} con {s:,.2f}")
        if c_noap>0: top_focus.append(f"Pagos banco no aplicados: {c_noap}")
        if c_pend>0: top_focus.append(f"Pendientes CxC: {c_pend}")
        if len(posibles_nc)>0: top_focus.append(f"Posibles NC/Retenciones sin cruzar: {len(posibles_nc)}")
        if not top_focus: top_focus.append("Sin focos críticos detectados.")

        sections = [
            ("RESUMEN EJECUTIVO", [f"• {x}" for x in resumen_doc]),
            ("HALLAZGOS RELEVANTES", [f"• {x}" for x in top_focus]),
            ("RECOMENDACIONES", [f"• {x}" for x in recs]),
            ("TRAZABILIDAD XLSX", ["• Ver 'cxc_bancos_hallazgos.xlsx' (todas las hojas)."])
        ]
        st.download_button("⬇️ Descargar reporte CxC vs Bancos (DOCX)",
                           docx_from_sections("CxC vs Bancos + Aging – Reporte de Auditoría", sections),
                           "reporte_cxc_bancos.docx",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ------------------------- Resumen gráfico (para 1–5) -------------------------
if opcion in PRUEBAS[:5] and sum(conteo_resultados.values()) > 0:
    st.subheader("📊 Resumen gráfico de pruebas CAAT")
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.barh(list(conteo_resultados.keys()), list(conteo_resultados.values()), color="steelblue")
    ax.set_xlabel("Cantidad"); ax.set_title("Resultados Detectados")
    st.pyplot(fig)

    st.subheader("🧠 Recomendaciones Automáticas")
    recomendaciones = [
        generar_recomendacion("Transacciones Conciliadas", conteo_resultados["Conciliadas"], 0,
                              "Conciliación correcta.", "Verifica registros coincidentes."),
        generar_recomendacion("Transacciones Faltantes", conteo_resultados["Faltantes en destino"], 2,
                              "Sin omisiones relevantes.", "Posibles errores u omisiones en registro."),
        generar_recomendacion("Transacciones Inesperadas", conteo_resultados["Inesperadas en destino"], 2,
                              "Sin ingresos inesperados.", "Revisar ingresos no respaldados por origen."),
        generar_recomendacion("Discrepancias de Valor", conteo_resultados["Discrepancias de valor"], 2,
                              "Fechas y montos alineados.", "Existen valores que no coinciden."),
        generar_recomendacion("Duplicados Internos", conteo_resultados["Duplicados"], 2,
                              "No hay duplicaciones.", "Registros repetidos requieren revisión.")
    ]
    for reco in recomendaciones:
        st.markdown(reco)

    st.subheader("🧾 Conclusión del Análisis")
    st.markdown(generar_conclusion_conteo(conteo_resultados))
